from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'LivePlay-Lab-产品作品集.docx'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def compact(paragraph, after=1, before=0):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = 1.08


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    compact(paragraph, after=3, before=6 if level == 1 else 3)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    compact(paragraph, after=2)
    return paragraph


def add_bullets(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.add_run(line)
        compact(paragraph, after=0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, 'DDEFE4')
        for run in cell.paragraphs[0].runs:
            run.bold = True
        compact(cell.paragraphs[0], after=0)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            compact(cells[idx].paragraphs[0], after=0)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
        row_properties = row._tr.get_or_add_trPr()
        cannot_split = OxmlElement('w:cantSplit')
        row_properties.append(cannot_split)
    return table


def configure(document):
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Microsoft YaHei'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(9.3)
    for style_name, size, color in [('Title', 22, '193B32'), ('Heading 1', 15, '1E6F55'), ('Heading 2', 11.5, '1E6F55')]:
        style = styles[style_name]
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    styles['List Bullet'].font.size = Pt(9.2)


def main():
    document = Document()
    configure(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('LivePlay Lab')
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor(25, 59, 50)
    compact(title, after=1)
    subtitle = document.add_paragraph('直播互动玩法生成与评测工作台｜独立产品作品集')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10.5)
    subtitle.runs[0].font.color.rgb = RGBColor(80, 102, 94)
    compact(subtitle, after=8)

    add_table(document, ['项目角色', '交付物', '版本状态'], [[
        '独立产品实践者：产品定义、规则设计、React 原型、评测设计、研究操作包',
        'MVP PRD、可运行 Demo、3 个模板、50 条评测、研究与录屏材料',
        'MVP 已实现；真实访谈、无引导测试、模型评测和 V2 数据待执行',
    ]], [2.35, 2.55, 2.3])

    add_heading(document, '01｜项目概述')
    add_body(document, 'LivePlay Lab 面向直播与校园活动组织者：把主题、目标、受众、时长和边界转成可执行玩法，并在浏览器中呈现可点击的互动微应用。产品进一步把生成质量放进 Promptfoo 评测，避免只凭“文案像不像”判断 Coding Agent。')
    add_body(document, '核心问题：玩法策划往往止于文本；AI 输出难以证明能跑通，也难以说明失败出在需求理解、代码状态还是自验证。')
    add_table(document, ['设计层', '产品回答'], [
        ['规则层', '三种模板将玩法拆为触发条件、参与动作、即时反馈和体验边界。'],
        ['体验层', '右侧浏览器微应用真实响应点击，并验证一次性提交与状态反馈。'],
        ['可靠性层', '50 条任务覆盖需求理解、代码生成、自验证、安全与体验。'],
        ['证据层', '研究记录和评测结果只有在真实执行后才回填，避免伪造样本或指标。'],
    ], [1.1, 6.1])

    add_heading(document, '02｜产品方案与流程')
    add_bullets(document, [
        '玩法工作台：填写 Brief，选择关键词拼图、限时真知题或任务接力赛；生成规则、口播、实现 Brief 和风险点。',
        '本地优先：未配置模型时依然生成结构化可运行草案；模型接口异常时回退到本地方案，保留用户输入。',
        '模型 API：支持 OpenAI-compatible chat-completions，通过本机代理调用；Key 只存在当前会话或本机 .env。',
        '模板与 Remix：将改过的规则保存为本机变体，刷新后可恢复，不存储密钥或真实用户资料。',
        '评测与研究：界面显示 50 条任务和 5+5 个研究位置；“已复核”和“待执行”严格区分。',
    ])
    add_table(document, ['主路径', '用户动作', '系统反馈'], [
        ['1. 定义问题', '填写主题、目标、受众、时长、边界', '对缺失主题/目标给出明确提示'],
        ['2. 选择规则', '选择一个模板', '显示规则摘要并同步微应用皮肤'],
        ['3. 生成方案', '选择本地或模型模式并生成', '输出结构化候选方案；模型失败回退'],
        ['4. 试跑玩法', '点击主互动', '锁定重复提交，显示积分或全场进度'],
        ['5. 沉淀证据', '保存 Remix、复核用例、执行研究', '本机保存规则；研究与评测只记录真实结果'],
    ], [1.15, 2.5, 3.55])

    add_heading(document, '03｜PRD 关键判断')
    add_table(document, ['决定', '理由', '验收方式'], [
        ['不做真实平台接入', 'MVP 先验证玩法和 Agent 产物闭环，避免账户、支付、隐私和平台依赖掩盖核心问题。', '代码和 UI 不出现平台授权、真实数据或品牌仿冒。'],
        ['默认本地生成', '作品集 Demo 需要无 Key、无网络也可展示。', '离线启动后可完成 Brief -> 方案 -> 微应用。'],
        ['模板先于自由生成', '先给用户可解释约束，降低黑盒感，也便于复用。', '三模板的规则、适用场景和预览状态都可见。'],
        ['失败分类进入产品', 'Agent 质量需要定位而非只看最终文本。', '50 条 Promptfoo 用例和人工复核字段可追溯。'],
    ], [1.35, 3.35, 2.5])

    add_heading(document, '04｜六周工作计划与工作日志')
    add_table(document, ['周次', '目标', '已产出 / 待执行', '反思'], [
        ['第 1 周', '访谈与竞品拆解', '研究包、招募与记录模板已完成；5 位真实访谈待执行。', '没有原始记录，不将假设写成洞察。'],
        ['第 2 周', 'MVP PRD、流程、三模板规则', 'PRD、Figma 复画说明、三套规则已完成。', '规则必须能映射到可观察的界面状态。'],
        ['第 3 周', 'Coding Agent 可运行 Demo 与模型 API', 'React Demo、本地生成、Remix、本机 API 代理已完成；真实 Key 调用待验证。', '本地回退让 Demo 不依赖外部服务。'],
        ['第 4 周', 'Promptfoo 50 条任务与失败分类', '配置与四类分类已完成；真实运行与人工复核待执行。', '没有原始报告，不报告通过率。'],
        ['第 5 周', '5 名新用户无引导测试', '任务卡、观察表、分析框架已完成；测试待执行。', '观察员提示会污染结果。'],
        ['第 6 周', 'V2、复盘、录屏、作品集', '作品集和录屏脚本已完成；V2 需由真实证据驱动。', '结论和证据必须一一对应。'],
    ], [0.75, 1.45, 3.2, 1.8])

    add_heading(document, '05｜工具使用与复现方法')
    add_table(document, ['工具', '作用', '使用边界'], [
        ['React + TypeScript + Vite', '搭建工作台、交互状态和响应式微应用。', '交付代码不连接真实平台。'],
        ['Node.js 本机服务', '代理 OpenAI-compatible 模型请求。', '不日志化 Key；.env 不入库。'],
        ['Promptfoo', '配置 50 条回归任务与 LLM rubric。', '真实运行后再输出结果和失败率。'],
        ['Markdown / Mermaid', 'PRD、流程、研究包和工作日志。', '文档中的待执行状态不替换为结论。'],
        ['浏览器自动化', '构建后检查主路径、桌面与移动布局。', '截图仅展示演示数据。'],
    ], [1.6, 3.0, 2.6])
    add_body(document, '复现：在项目根目录运行 npm install、npm run dev；执行 npm run test、npm run lint、npm run build。模型调用需要另开终端运行 npm run server，并在界面会话或本机 .env 填写兼容端点与 Key。', '复现：')

    add_heading(document, '06｜研究与评测计划（未执行部分）')
    add_body(document, '第 1 周招募 5 名有直播或活动组织经验的人，使用 30 分钟半结构访谈理解玩法如何从想法走到现场。第 5 周招募 5 名新用户，在不提示的情况下完成“选择模板 -> 生成 -> 预览 -> 保存 Remix”。每次观察要区分原话、事实和研究者解释。')
    add_body(document, '评测使用 50 条 Promptfoo 用例。每条需要记录模型、提示版本、原始输出、人工复核、失败编码和修复假设。无论是用户研究还是模型评测，执行前都不得在作品集中填“完成率”“提升率”或虚构反馈。')

    add_heading(document, '07｜下一步与面试表达')
    add_bullets(document, [
        '先完成 5+5 真实研究，按严重度和频次形成 V2 backlog。',
        '使用目标模型跑完 50 条评测，优先处理安全与主路径严重失败，并保留前后对比。',
        '录制 5 分钟演示：展示 Brief、模板、可点击状态、失败回退、评测和研究边界。',
        '面试时从“如何把 AI 输出变成可验证产品”讲起：约束输入、规则化生成、可运行原型、失败分类、真实研究闭环。',
    ])
    add_body(document, '项目材料：README、PRD、Figma 流程、研究包、评测方案、六周工作日志、操作手册、演示脚本均位于仓库 docs/ 目录。')

    document.save(OUT)
    print(OUT)


if __name__ == '__main__':
    main()
